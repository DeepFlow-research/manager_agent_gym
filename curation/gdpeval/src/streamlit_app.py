"""Streamlit app for reviewing staged rubrics with gates and thresholds."""

import streamlit as st  # type: ignore[import-untyped]
import pandas as pd
import json
from pathlib import Path
from datetime import datetime
import fitz  # PyMuPDF for PDF preview
from docx import Document  # python-docx for Word documents
from PIL import Image
import io

# Page config
st.set_page_config(
    page_title="GDPEval Staged Rubric Annotation",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Paths
BASE_DIR = Path(__file__).parent.parent
STAGED_RUBRICS_FILE = (
    BASE_DIR / "data" / "generated" / "staged_v1" / "staged_rubrics.jsonl"
)
FEEDBACK_FILE = BASE_DIR / "data" / "feedback" / "staged_feedback.jsonl"
GDPEVAL_FILE = BASE_DIR / "data" / "raw" / "gdpeval.parquet"
REF_FILES_DIR = BASE_DIR / "data" / "raw" / "reference_files"


# Load data
@st.cache_data
def load_rubrics():
    """Load generated staged rubrics."""
    if not STAGED_RUBRICS_FILE.exists():
        return pd.DataFrame()
    rubrics = []
    with open(STAGED_RUBRICS_FILE, "r") as f:
        for line in f:
            rubrics.append(json.loads(line))
    return pd.DataFrame(rubrics)


@st.cache_data
def load_gdpeval():
    """Load GDPEval dataset."""
    if not GDPEVAL_FILE.exists():
        return pd.DataFrame()
    return pd.read_parquet(GDPEVAL_FILE)


@st.cache_data
def load_feedback() -> pd.DataFrame:
    """Load existing feedback."""
    if not FEEDBACK_FILE.exists():
        return pd.DataFrame(
            columns=["task_id", "timestamp", "status", "comments", "updated_rubric"]
        )  # type: ignore[call-overload]

    feedback = []
    with open(FEEDBACK_FILE, "r") as f:
        for line in f:
            feedback.append(json.loads(line))
    return pd.DataFrame(feedback)


def save_feedback(task_id: str, status: str, comments: str):
    """Save feedback for a task."""
    FEEDBACK_FILE.parent.mkdir(parents=True, exist_ok=True)

    feedback_entry = {
        "task_id": task_id,
        "timestamp": datetime.now().isoformat(),
        "status": status,
        "comments": comments,
    }

    with open(FEEDBACK_FILE, "a") as f:
        f.write(json.dumps(feedback_entry) + "\n")


def render_file_preview(file_path: Path):
    """Render preview for various file types."""
    if not file_path.exists():
        st.warning(f"File not found: {file_path.name}")
        return

    file_ext = file_path.suffix.lower()

    try:
        # PDF Preview
        if file_ext == ".pdf":
            pdf_doc = fitz.open(file_path)
            st.write(f"📄 **{file_path.name}** ({len(pdf_doc)} pages)")

            # Show first 3 pages as images
            for page_num in range(min(3, len(pdf_doc))):
                page = pdf_doc[page_num]
                pix = page.get_pixmap(dpi=150)
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                st.image(img, caption=f"Page {page_num + 1}", use_container_width=True)
                if page_num < min(3, len(pdf_doc)) - 1:
                    st.divider()

            if len(pdf_doc) > 3:
                st.info(f"... {len(pdf_doc) - 3} more pages")

        # Word Document Preview
        elif file_ext in [".docx", ".doc"]:
            doc = Document(str(file_path))
            st.write(f"📝 **{file_path.name}** ({len(doc.paragraphs)} paragraphs)")

            # Show first few paragraphs
            preview_text = "\n\n".join(
                [p.text for p in doc.paragraphs[:10] if p.text.strip()]
            )
            st.text_area("Preview", preview_text, height=300, disabled=True)

            if len(doc.paragraphs) > 10:
                st.info(f"... {len(doc.paragraphs) - 10} more paragraphs")

        # Excel Preview
        elif file_ext in [".xlsx", ".xls", ".xlsm"]:
            excel_data = pd.ExcelFile(file_path)
            st.write(f"📊 **{file_path.name}** ({len(excel_data.sheet_names)} sheets)")

            for sheet_name in excel_data.sheet_names[:3]:
                st.write(f"**Sheet: {sheet_name}**")
                df = pd.read_excel(file_path, sheet_name=sheet_name, nrows=10)
                st.dataframe(df, use_container_width=True)
                if len(excel_data.sheet_names) > 1:
                    st.divider()

            if len(excel_data.sheet_names) > 3:
                st.info(f"... {len(excel_data.sheet_names) - 3} more sheets")

        # CSV Preview
        elif file_ext == ".csv":
            df = pd.read_csv(file_path, nrows=20)
            st.write(f"📈 **{file_path.name}** ({len(df)} rows shown)")
            st.dataframe(df, use_container_width=True)

        # Image Preview
        elif file_ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp"]:
            img = Image.open(file_path)
            st.image(img, caption=file_path.name, use_container_width=True)

        # Text File Preview
        elif file_ext in [
            ".txt",
            ".md",
            ".json",
            ".xml",
            ".html",
            ".css",
            ".js",
            ".py",
        ]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read(5000)  # First 5000 chars
            st.code(
                content,
                language=file_ext[1:] if file_ext in [".py", ".js", ".html"] else None,
            )

            if len(content) == 5000:
                st.info("Preview truncated...")

        else:
            st.info(f"Preview not available for {file_ext} files")

        # Always show file path for opening
        st.caption(f"📁 Open in default app: `{file_path}`")

    except Exception as e:
        st.error(f"Error previewing file: {str(e)}")


def render_stage(stage_data: dict, stage_num: int):
    """Render a single evaluation stage."""
    is_gate = stage_data.get("is_required", False)

    # Stage header with visual indicators
    gate_badge = "🚪 GATE" if is_gate else "📋 STAGE"
    failure_action = stage_data.get("on_failure_action", "continue")

    with st.container(border=True):
        col1, col2, col3 = st.columns([2, 1, 1])

        with col1:
            st.subheader(f"{gate_badge} Stage {stage_num}: {stage_data['name']}")
            st.caption(stage_data.get("description", ""))

        with col2:
            st.metric("Max Points", f"{stage_data['max_points']:.1f}")

        with col3:
            if is_gate:
                threshold = stage_data.get("min_score_to_pass", 1.0)
                st.metric("Pass Threshold", f"{threshold * 100:.0f}%")

        # Show failure action for gates
        if is_gate:
            if failure_action == "zero_category":
                st.error("⚠️ **If this gate fails → Entire category gets 0 points**")
            elif failure_action == "skip_remaining":
                st.warning("⚠️ **If this gate fails → Skip remaining stages**")

        # Render rules
        st.write("**Rules:**")
        for idx, rule in enumerate(stage_data.get("rules", [])):
            rule_type = rule.get("type", "unknown")
            rule_icon = (
                "💻"
                if rule_type == "code"
                else "🤖"
                if rule_type == "llm_judge"
                else "❓"
            )

            with st.expander(
                f"{rule_icon} {rule['name']} (Weight: {rule.get('weight', 1.0)})"
            ):
                st.write(f"**Description:** {rule.get('description', 'N/A')}")

                if rule_type == "code":
                    st.code(rule.get("code", "No code provided"), language="python")

                elif rule_type == "llm_judge":
                    st.write("**Judge Prompt:**")
                    st.info(rule.get("judge_prompt", "No prompt provided"))
                    if "expectation" in rule:
                        st.write("**Expectation:**")
                        st.success(rule["expectation"])


# Main app
st.title("🎯 GDPEval Staged Rubric Review")

# Load data
rubrics_df = load_rubrics()
gdpeval_df = load_gdpeval()
feedback_df = load_feedback()

# Sidebar - Progress & Navigation
with st.sidebar:
    st.header("📊 Progress")

    if len(rubrics_df) == 0:
        st.warning("No staged rubrics generated yet!")
        st.info(
            "Run: `python -m curation.gdpeval.src.generate_staged_rubrics --limit 10`"
        )
        st.stop()

    # Calculate stats
    total_tasks = len(gdpeval_df)
    generated = len(rubrics_df)
    annotated = len(feedback_df)

    col1, col2 = st.columns(2)
    col1.metric("Generated", f"{generated}/{total_tasks}")
    col2.metric("Annotated", annotated)

    st.progress(annotated / max(generated, 1))

    st.divider()

    # Filters
    st.header("🔍 Filters")

    # Status filter
    annotated_task_ids = (
        set(feedback_df["task_id"].tolist()) if len(feedback_df) > 0 else set()
    )

    status_options = {
        "All": None,
        "Not Annotated": lambda df: ~df["task_id"].isin(annotated_task_ids),
        "Annotated": lambda df: df["task_id"].isin(annotated_task_ids),
    }
    status_filter = st.selectbox("Status", list(status_options.keys()))

    # Sector filter
    if len(gdpeval_df) > 0:
        sectors = ["All"] + sorted(gdpeval_df["sector"].unique().tolist())
        sector_filter = st.selectbox("Sector", sectors)
    else:
        sector_filter = "All"

# Apply filters
filtered_df = rubrics_df.copy()

if status_filter != "All":
    filter_fn = status_options[status_filter]
    filtered_df = filtered_df[filter_fn(filtered_df)]

if sector_filter != "All" and len(gdpeval_df) > 0:
    sector_task_ids = gdpeval_df[gdpeval_df["sector"] == sector_filter][
        "task_id"
    ].tolist()
    # Type checker gets confused about DataFrame vs ndarray, so we assert the type
    assert isinstance(filtered_df, pd.DataFrame)
    filtered_df = filtered_df[filtered_df["task_id"].isin(sector_task_ids)]

# Task selection
if len(filtered_df) == 0:
    st.info("No tasks match the selected filters.")
    st.stop()

# Task selector
task_ids = filtered_df["task_id"].tolist()
current_idx = st.sidebar.number_input(
    "Task Index",
    min_value=0,
    max_value=len(task_ids) - 1,
    value=0,
)

current_task_id = task_ids[current_idx]
st.sidebar.caption(f"Task ID: {current_task_id[:8]}...")

# Navigation buttons
nav_col1, nav_col2 = st.sidebar.columns(2)
if nav_col1.button("⬅️ Previous", disabled=(current_idx == 0)):
    st.session_state.task_idx = current_idx - 1
if nav_col2.button("Next ➡️", disabled=(current_idx == len(task_ids) - 1)):
    st.session_state.task_idx = current_idx + 1

# Main content
assert isinstance(filtered_df, pd.DataFrame)  # Type hint for checker
rubric_row = filtered_df[filtered_df["task_id"] == current_task_id].iloc[0]
task_row = (
    gdpeval_df[gdpeval_df["task_id"] == current_task_id].iloc[0]
    if len(gdpeval_df) > 0
    else None
)

# Task metadata
st.header("📋 Task Information")

if task_row is not None:
    col1, col2, col3 = st.columns(3)
    col1.metric("Sector", task_row["sector"])
    col2.metric("Occupation", task_row["occupation"])
    col3.metric("Reference Files", len(task_row.get("reference_files", [])))

    st.subheader("Task Description")
    st.write(task_row["prompt"])
else:
    st.info(f"Task ID: {current_task_id}")

st.divider()

# Reference files
if task_row is not None and len(task_row.get("reference_files", [])) > 0:
    st.header("📁 Reference Files")

    ref_files = task_row.get("reference_files", [])
    task_ref_dir = REF_FILES_DIR / current_task_id

    if task_ref_dir.exists():
        for ref_file in ref_files:
            file_path = task_ref_dir / Path(ref_file).name
            if file_path.exists():
                with st.expander(f"📄 {file_path.name}", expanded=False):
                    render_file_preview(file_path)
    else:
        st.info("Reference files not downloaded yet")

    st.divider()

# Staged Rubric Display
st.header("🎯 Staged Evaluation Rubric")

rubric_data = rubric_row["rubric"]

# Rubric overview
col1, col2 = st.columns([2, 1])
with col1:
    st.subheader(rubric_data["category_name"])
    if rubric_data.get("rationale"):
        st.info(f"**Rationale:** {rubric_data['rationale']}")

with col2:
    st.metric("Total Max Score", f"{rubric_data['max_total_score']:.1f}")
    st.metric("Number of Stages", len(rubric_data["stages"]))

st.divider()

# Render each stage
for idx, stage in enumerate(rubric_data["stages"], 1):
    render_stage(stage, idx)

st.divider()

# Feedback section
st.header("💬 Feedback")

existing_feedback = feedback_df[feedback_df["task_id"] == current_task_id]
if len(existing_feedback) > 0:
    latest = existing_feedback.iloc[-1]
    st.success(f"✅ Annotated on {latest['timestamp']}")
    st.write(f"**Status:** {latest['status']}")
    st.write(f"**Comments:** {latest['comments']}")
    st.divider()

# New feedback form
with st.form("feedback_form"):
    st.subheader("Add Feedback")

    status = st.selectbox(
        "Status",
        ["Approve", "Needs Revision", "Reject"],
        help="Overall assessment of this staged rubric",
    )

    comments = st.text_area(
        "Comments",
        placeholder="Share your thoughts on:\n- Are the stages well-designed?\n- Are gates appropriate?\n- Are the rules clear and verifiable?\n- Any improvements needed?",
        height=150,
    )

    submit = st.form_submit_button("💾 Save Feedback")

    if submit:
        if not comments.strip():
            st.error("Please provide comments!")
        else:
            save_feedback(current_task_id, status, comments)
            st.success("Feedback saved!")
            st.rerun()

# Footer
st.divider()
st.caption(
    f"Task {current_idx + 1} of {len(filtered_df)} (filtered) | {len(rubrics_df)} total staged rubrics generated"
)
