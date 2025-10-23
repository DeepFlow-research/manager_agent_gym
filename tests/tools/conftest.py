"""Comprehensive pytest configuration and fixtures for tool testing.

Provides:
- Pytest markers for external dependencies
- API key fixtures with environment checking
- Resource management fixtures
- Sample file generation fixtures
- Test utilities and helpers
"""

import os
from pathlib import Path
from typing import Generator

import pytest


# ============================================================================
# PYTEST MARKERS
# ============================================================================


def pytest_configure(config: pytest.Config) -> None:
    """Register custom pytest markers."""
    config.addinivalue_line(
        "markers",
        "requires_e2b: mark test as requiring E2B_API_KEY environment variable",
    )
    config.addinivalue_line(
        "markers",
        "requires_cohere: mark test as requiring COHERE_API_KEY environment variable",
    )
    config.addinivalue_line(
        "markers",
        "requires_exa: mark test as requiring EXA_API_KEY environment variable",
    )
    config.addinivalue_line(
        "markers",
        "requires_tesseract: mark test as requiring Tesseract OCR installed on system",
    )
    config.addinivalue_line(
        "markers",
        "requires_libreoffice: mark test as requiring LibreOffice (soffice) installed on system",
    )
    config.addinivalue_line(
        "markers",
        "requires_rag: mark test as requiring RAG dependencies (rank-bm25)",
    )


# ============================================================================
# DEPENDENCY CHECKING
# ============================================================================


@pytest.fixture(autouse=True)
def check_rag_dependencies(request):
    """Auto-skip tests marked with requires_rag if dependencies not installed."""
    marker = request.node.get_closest_marker("requires_rag")
    if marker:
        try:
            import rank_bm25
        except ImportError:
            pytest.skip("RAG dependencies not installed (rank-bm25)")


# ============================================================================
# API KEY FIXTURES
# ============================================================================


@pytest.fixture
def e2b_api_key() -> Generator[str, None, None]:
    """E2B API key fixture - skips test if not available."""
    api_key = os.getenv("E2B_API_KEY")
    if not api_key:
        pytest.skip("E2B_API_KEY not set")
    yield api_key


@pytest.fixture
def cohere_api_key() -> Generator[str, None, None]:
    """Cohere API key fixture - skips test if not available."""
    api_key = os.getenv("COHERE_API_KEY")
    if not api_key:
        pytest.skip("COHERE_API_KEY not set")
    yield api_key


@pytest.fixture
def exa_api_key() -> Generator[str, None, None]:
    """Exa API key fixture - skips test if not available."""
    api_key = os.getenv("EXA_API_KEY")
    if not api_key:
        pytest.skip("EXA_API_KEY not set")
    yield api_key


# ============================================================================
# RESOURCE MANAGEMENT FIXTURES
# ============================================================================


@pytest.fixture
def resource_manager(tmp_path: Path):
    """ResourceFileManager with temp directory."""
    from manager_agent_gym.core.workflow.resource_storage import ResourceFileManager

    return ResourceFileManager()


@pytest.fixture
def tmp_directory(tmp_path: Path) -> Path:
    """Temporary directory for test files."""
    return tmp_path


# ============================================================================
# SAMPLE FILE FIXTURES
# ============================================================================


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """Generate a sample PDF file with text and tables."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
        )
        from reportlab.lib.units import inch
        from reportlab.lib import colors
    except ImportError:
        pytest.skip("reportlab not installed")

    pdf_path = tmp_path / "sample.pdf"

    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Add title
    story.append(Paragraph("Test Document", styles["Title"]))
    story.append(Spacer(1, 0.3 * inch))

    # Add body text
    story.append(
        Paragraph(
            "This is a test PDF document with sample content for testing PDF reading tools.",
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 0.2 * inch))

    # Add table
    table_data = [
        ["Name", "Age", "City"],
        ["Alice", "30", "New York"],
        ["Bob", "25", "London"],
        ["Charlie", "35", "Paris"],
    ]
    table = Table(table_data)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 14),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]
        )
    )
    story.append(table)

    doc.build(story)
    return pdf_path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Generate a sample DOCX file."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        pytest.skip("python-docx not installed")

    docx_path = tmp_path / "sample.docx"

    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test DOCX document with sample content.")
    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("Content for section 1.")

    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Header 1"
    hdr_cells[1].text = "Header 2"
    hdr_cells[2].text = "Header 3"

    row_cells = table.rows[1].cells
    row_cells[0].text = "Data 1"
    row_cells[1].text = "Data 2"
    row_cells[2].text = "Data 3"

    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def sample_csv(tmp_path: Path) -> Path:
    """Generate a sample CSV file."""
    try:
        import pandas as pd
    except ImportError:
        pytest.skip("pandas not installed")

    csv_path = tmp_path / "sample.csv"

    data = {
        "Name": ["Alice", "Bob", "Charlie", "Diana"],
        "Age": [30, 25, 35, 28],
        "City": ["New York", "London", "Paris", "Tokyo"],
        "Salary": [75000, 65000, 85000, 70000],
    }

    df = pd.DataFrame(data)
    df.to_csv(str(csv_path), index=False)
    return csv_path


@pytest.fixture
def sample_excel(tmp_path: Path) -> Path:
    """Generate a sample Excel file."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill
    except ImportError:
        pytest.skip("openpyxl not installed")

    excel_path = tmp_path / "sample.xlsx"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"

    # Add headers
    headers = ["Name", "Age", "City", "Salary"]
    ws.append(headers)

    # Style headers
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(
            start_color="4472C4", end_color="4472C4", fill_type="solid"
        )

    # Add data rows
    data_rows = [
        ["Alice", 30, "New York", 75000],
        ["Bob", 25, "London", 65000],
        ["Charlie", 35, "Paris", 85000],
        ["Diana", 28, "Tokyo", 70000],
    ]

    for row in data_rows:
        ws.append(row)

    wb.save(str(excel_path))
    return excel_path


@pytest.fixture
def sample_markdown(tmp_path: Path) -> Path:
    """Generate a sample Markdown file."""
    md_path = tmp_path / "sample.md"

    content = """# Test Document

This is a test markdown document.

## Section 1

Some content here.

- Item 1
- Item 2
- Item 3

## Section 2

More content.

```python
def hello():
    print("Hello, World!")
```

| Name | Age |
|------|-----|
| Alice | 30 |
| Bob | 25 |
"""

    md_path.write_text(content)
    return md_path


@pytest.fixture
def sample_image_with_text(tmp_path: Path) -> Path:
    """Generate a sample image with text for OCR testing."""
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        pytest.skip("Pillow not installed")

    image_path = tmp_path / "sample_with_text.png"

    # Create image with text
    image = Image.new("RGB", (400, 200), color="white")
    draw = ImageDraw.Draw(image)

    # Try to use a default font, fallback to default if not available
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
    except (OSError, IOError):
        font = ImageFont.load_default()

    text = "Hello World\nThis is a test image\nfor OCR testing"
    draw.multiline_text((50, 50), text, fill="black", font=font)

    image.save(str(image_path))
    return image_path


@pytest.fixture
def sample_text_file(tmp_path: Path) -> Path:
    """Generate a sample text file."""
    text_path = tmp_path / "sample.txt"

    content = """Sample Text Document

This is a test text file with multiple paragraphs.

First section:
- Point 1
- Point 2
- Point 3

Second section:
Some more content here to make it more realistic.
"""

    text_path.write_text(content)
    return text_path


# ============================================================================
# HELPER FIXTURES
# ============================================================================


@pytest.fixture
def cleanup_temp_files(tmp_path: Path) -> Generator[None, None, None]:
    """Cleanup fixture for temporary files."""
    yield
    # Cleanup is handled automatically by tmp_path


@pytest.fixture(scope="session")
def tesseract_available() -> bool:
    """Check if Tesseract OCR is installed."""
    import shutil

    return shutil.which("tesseract") is not None


@pytest.fixture(scope="session")
def libreoffice_available() -> bool:
    """Check if LibreOffice is installed."""
    import shutil

    return shutil.which("soffice") is not None


# ============================================================================
# PYTEST HOOKS
# ============================================================================


def pytest_runtest_setup(item: pytest.Item) -> None:
    """Skip tests based on markers and environment."""
    # Check requires_tesseract marker
    if "requires_tesseract" in item.keywords:
        import shutil

        if not shutil.which("tesseract"):
            pytest.skip("Tesseract OCR not installed")

    # Check requires_libreoffice marker
    if "requires_libreoffice" in item.keywords:
        import shutil

        if not shutil.which("soffice"):
            pytest.skip("LibreOffice not installed")
