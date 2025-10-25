"""Document processing tools (PDF, DOCX, Markdown) - two-layer architecture.

Layer 1: Core functions (_*) - pure business logic, testable, returns typed results
Layer 2: OpenAI tool wrappers - thin adapters for OpenAI SDK, handle JSON serialization
"""

import json
import subprocess
import tempfile
from pathlib import Path
from typing import TYPE_CHECKING, Any

import pdfplumber
from agents import Tool, function_tool
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PyPDF2 import PdfMerger
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

if TYPE_CHECKING:
    from manager_agent_gym.core.workflow.resource_storage import ResourceFileManager


# ============================================================================
# LAYER 1: PDF OPERATIONS (Core Business Logic)
# ============================================================================


async def _read_pdf(file_path: str) -> dict[str, Any]:
    """Extract text from PDF file with page markers."""
    try:
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            return {"success": False, "error": f"PDF file not found at {file_path}"}

        extracted_text = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text:
                    extracted_text.append(f"--- Page {page_num} ---\n{text}")

        if not extracted_text:
            return {
                "success": False,
                "error": "No text could be extracted from the PDF",
            }

        return {"success": True, "text": "\n\n".join(extracted_text)}

    except Exception as e:
        return {"success": False, "error": f"Error reading PDF: {str(e)}"}


async def _extract_pdf_tables(
    file_path: str, page_number: int | None = None
) -> dict[str, Any]:
    """Extract tables from PDF file."""
    try:
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            return {"success": False, "error": f"PDF file not found at {file_path}"}

        all_tables = []
        with pdfplumber.open(file_path) as pdf:
            pages_to_process = (
                [pdf.pages[page_number - 1]] if page_number else pdf.pages
            )

            for page_idx, page in enumerate(pages_to_process, start=1):
                actual_page_num = page_number if page_number else page_idx
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables, start=1):
                    all_tables.append(
                        {
                            "page": actual_page_num,
                            "table_index": table_idx,
                            "data": table,
                        }
                    )

        if not all_tables:
            return {"success": True, "tables": [], "message": "No tables found"}

        return {"success": True, "tables": all_tables}

    except Exception as e:
        return {"success": False, "error": f"Error extracting tables: {str(e)}"}


async def _merge_pdfs(pdf_paths: list[str], output_path: str) -> dict[str, Any]:
    """Merge multiple PDF files."""
    try:
        if not pdf_paths:
            return {"success": False, "error": "No PDF files provided"}

        for pdf_path in pdf_paths:
            if not Path(pdf_path).exists():
                return {"success": False, "error": f"PDF not found: {pdf_path}"}

        merger = PdfMerger()
        for pdf_path in pdf_paths:
            merger.append(pdf_path)

        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)

        merger.write(output_path)
        merger.close()

        return {
            "success": True,
            "output_path": str(output_path_obj.absolute()),
            "merged_count": len(pdf_paths),
        }

    except Exception as e:
        return {"success": False, "error": f"Error merging PDFs: {str(e)}"}


async def _create_pdf(
    text_content: str, output_path: str, title: str | None = None
) -> dict[str, Any]:
    """Create PDF from text content."""
    try:
        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)

        doc = SimpleDocTemplate(str(output_path_obj), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        if title:
            story.append(Paragraph(title, styles["Title"]))
            story.append(Spacer(1, 0.2 * inch))

        body_style = ParagraphStyle(
            "CustomBody",
            parent=styles["BodyText"],
            alignment=TA_JUSTIFY,
            fontSize=11,
            leading=14,
        )

        for para_text in text_content.split("\n\n"):
            if para_text.strip():
                para_text = para_text.replace("\n", "<br/>")
                story.append(Paragraph(para_text, body_style))
                story.append(Spacer(1, 0.1 * inch))

        doc.build(story)

        return {
            "success": True,
            "output_path": str(output_path_obj.absolute()),
            "file_size": output_path_obj.stat().st_size,
        }

    except Exception as e:
        return {"success": False, "error": f"Error creating PDF: {str(e)}"}


# ============================================================================
# LAYER 1: DOCX OPERATIONS
# ============================================================================


async def _read_docx(file_path: str) -> dict[str, Any]:
    """Extract text from DOCX file."""
    try:
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            return {"success": False, "error": f"DOCX file not found at {file_path}"}

        doc = Document(file_path)

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        if not paragraphs:
            return {"success": False, "error": "No text could be extracted"}

        return {"success": True, "text": "\n\n".join(paragraphs)}

    except Exception as e:
        return {"success": False, "error": f"Error reading DOCX: {str(e)}"}


async def _create_docx(
    content: str,
    output_path: str,
    title: str | None = None,
    template_style: str = "normal",
) -> dict[str, Any]:
    """Create DOCX file from text content."""
    try:
        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        sections = doc.sections
        for section in sections:
            if template_style == "formal":
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            elif template_style == "memo":
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
            else:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)

        if title:
            title_para = doc.add_paragraph(title)
            title_para.style = "Title"
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for para_text in content.split("\n\n"):
            if para_text.strip():
                if para_text.startswith("# "):
                    para = doc.add_paragraph(para_text[2:])
                    para.style = "Heading 1"  # type: ignore[assignment]
                elif para_text.startswith("## "):
                    para = doc.add_paragraph(para_text[3:])
                    para.style = "Heading 2"  # type: ignore[assignment]
                elif para_text.startswith("### "):
                    para = doc.add_paragraph(para_text[4:])
                    para.style = "Heading 3"  # type: ignore[assignment]
                else:
                    para = doc.add_paragraph(para_text)
                    if template_style == "formal":
                        for run in para.runs:
                            run.font.size = Pt(12)
                    elif template_style == "memo":
                        for run in para.runs:
                            run.font.size = Pt(11)

        doc.save(str(output_path_obj))

        return {
            "success": True,
            "output_path": str(output_path_obj.absolute()),
            "file_size": output_path_obj.stat().st_size,
        }

    except Exception as e:
        return {"success": False, "error": f"Error creating DOCX: {str(e)}"}


async def _add_table_to_docx(
    file_path: str, table_data: list[list[str]], header_row: bool = True
) -> dict[str, Any]:
    """Add table to existing DOCX file."""
    try:
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            return {"success": False, "error": "DOCX file not found"}

        if not table_data or not table_data[0]:
            return {"success": False, "error": "Table data is empty"}

        doc = Document(file_path)

        num_rows = len(table_data)
        num_cols = len(table_data[0])
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Light Grid Accent 1"

        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_value in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(cell_value)

                if row_idx == 0 and header_row:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(11)

        doc.save(str(file_path_obj))

        return {"success": True, "output_path": str(file_path_obj.absolute())}

    except Exception as e:
        return {"success": False, "error": f"Error adding table: {str(e)}"}


async def _extract_docx_structure(file_path: str) -> dict[str, Any]:
    """Extract document structure from DOCX."""
    try:
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            return {"success": False, "error": "DOCX file not found"}

        doc = Document(file_path)
        structure = []

        for element in doc.element.body:
            if element.tag.endswith("p"):
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                if para and para.text.strip():
                    style_name = para.style.name if para.style else "Normal"
                    structure.append(
                        {
                            "type": "paragraph",
                            "style": style_name,
                            "text": para.text[:100]
                            + ("..." if len(para.text) > 100 else ""),
                        }
                    )
            elif element.tag.endswith("tbl"):
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break
                if table:
                    structure.append(
                        {
                            "type": "table",
                            "rows": len(table.rows),
                            "columns": len(table.columns),
                        }
                    )

        return {"success": True, "structure": structure}

    except Exception as e:
        return {"success": False, "error": f"Error extracting structure: {str(e)}"}


# ============================================================================
# LAYER 1: DOCUMENT CONVERTERS
# ============================================================================


async def _convert_docx_to_pdf(docx_path: str, output_path: str) -> dict[str, Any]:
    """Convert DOCX to PDF using LibreOffice."""
    try:
        docx_path_obj = Path(docx_path)
        if not docx_path_obj.exists():
            return {"success": False, "error": "DOCX file not found"}

        # Validate that the input is actually a .docx file
        if docx_path_obj.suffix.lower() not in [".docx", ".doc"]:
            return {
                "success": False,
                "error": f"Invalid file type '{docx_path_obj.suffix}'. This tool only converts Word documents (.docx, .doc). Use 'convert_markdown_to_docx' first if you have a Markdown file.",
            }

        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)

        try:
            output_dir = output_path_obj.parent
            result = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_dir),
                    str(docx_path_obj),
                ],
                capture_output=True,
                text=True,
                timeout=30,
            )

            if result.returncode == 0:
                generated_pdf = output_dir / f"{docx_path_obj.stem}.pdf"
                if generated_pdf.exists() and generated_pdf != output_path_obj:
                    generated_pdf.rename(output_path_obj)
                return {
                    "success": True,
                    "output_path": str(output_path_obj.absolute()),
                }
            else:
                return {
                    "success": False,
                    "error": f"LibreOffice error: {result.stderr}",
                }

        except FileNotFoundError:
            return {
                "success": False,
                "error": "LibreOffice (soffice) not found. Please install LibreOffice.",
            }
        except subprocess.TimeoutExpired:
            return {"success": False, "error": "Conversion timed out"}

    except Exception as e:
        return {"success": False, "error": f"Error converting DOCX to PDF: {str(e)}"}


async def _convert_markdown_to_docx(
    markdown_content: str, output_path: str
) -> dict[str, Any]:
    """Convert Markdown to DOCX."""
    try:
        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        lines = markdown_content.split("\n")
        in_code_block = False
        code_lines: list[str] = []

        for line in lines:
            if line.strip().startswith("```"):
                if in_code_block:
                    if code_lines:
                        para = doc.add_paragraph("\n".join(code_lines))
                        para.style = "Code"  # type: ignore[assignment]
                        code_lines = []
                in_code_block = not in_code_block
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            if line.startswith("# "):
                para = doc.add_paragraph(line[2:])
                para.style = "Heading 1"  # type: ignore[assignment]
            elif line.startswith("## "):
                para = doc.add_paragraph(line[3:])
                para.style = "Heading 2"  # type: ignore[assignment]
            elif line.startswith("### "):
                para = doc.add_paragraph(line[4:])
                para.style = "Heading 3"  # type: ignore[assignment]
            elif line.strip().startswith("- ") or line.strip().startswith("* "):
                para = doc.add_paragraph(line.strip()[2:], style="List Bullet")
            elif line.strip() and line.strip()[0].isdigit() and ". " in line:
                content = line.strip().split(". ", 1)[1]
                para = doc.add_paragraph(content, style="List Number")
            elif line.strip():
                para = doc.add_paragraph(line)
            else:
                doc.add_paragraph()

        doc.save(str(output_path_obj))

        return {
            "success": True,
            "output_path": str(output_path_obj.absolute()),
            "file_size": output_path_obj.stat().st_size,
        }

    except Exception as e:
        return {"success": False, "error": f"Error converting Markdown: {str(e)}"}


async def _convert_text_to_docx(
    text_content: str, output_path: str, font_size: int = 11
) -> dict[str, Any]:
    """Convert plain text to DOCX."""
    try:
        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        for para_text in text_content.split("\n\n"):
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                for run in para.runs:
                    run.font.size = Pt(font_size)

        doc.save(str(output_path_obj))

        return {
            "success": True,
            "output_path": str(output_path_obj.absolute()),
            "file_size": output_path_obj.stat().st_size,
        }

    except Exception as e:
        return {"success": False, "error": f"Error converting text: {str(e)}"}


# ============================================================================
# LAYER 1: MARKDOWN
# ============================================================================


async def _save_markdown(content: str, file_name: str) -> dict[str, Any]:
    """Save text as Markdown file."""
    try:
        if not file_name.endswith(".md"):
            file_name = f"{file_name}.md"

        temp_dir = Path(tempfile.gettempdir()) / "workflow_outputs"
        temp_dir.mkdir(parents=True, exist_ok=True)

        file_path = temp_dir / file_name
        file_path.write_text(content, encoding="utf-8")
        file_size = file_path.stat().st_size

        lines = content.split("\n")
        num_lines = len(lines)
        num_headings = sum(1 for line in lines if line.strip().startswith("#"))
        word_count = len(content.split())

        preview = content[:500] + "..." if len(content) > 500 else content

        return {
            "success": True,
            "file_path": str(file_path.absolute()),
            "file_name": file_name,
            "mime_type": "text/markdown",
            "size_bytes": file_size,
            "content_summary": {
                "num_lines": num_lines,
                "num_headings": num_headings,
                "word_count": word_count,
                "preview": preview,
            },
        }

    except Exception as e:
        return {"success": False, "error": f"Failed to save markdown: {str(e)}"}


# ============================================================================
# LAYER 2: OPENAI TOOL WRAPPERS
# ============================================================================


def create_documents_tools(resource_manager: "ResourceFileManager") -> list[Tool]:
    """Create document processing tools for OpenAI SDK."""
    from manager_agent_gym.core.workflow.context import AgentExecutionContext
    from agents import RunContextWrapper

    @function_tool
    async def read_pdf(file_path: str) -> str:
        """
        Extract text content from a PDF file with page markers for easy navigation.

        This tool reads PDF files and extracts all text content, inserting page markers
        (e.g., "--- Page 1 ---") to help you identify which page text came from. Perfect
        for analyzing reports, contracts, research papers, or any PDF documents.

        Parameters:
            file_path (str):
                The absolute or relative path to the PDF file to read.
                Example: "/path/to/document.pdf" or "reports/quarterly_report.pdf"

        Returns:
            str:
                The full text content of the PDF with page markers, or an error message if
                the file cannot be read or contains no extractable text.

        Usage:
            Call this tool when you need to read and analyze PDF documents. Use it for
            extracting information, summarizing content, or searching for specific data
            within PDF files.
        """
        result = await _read_pdf(file_path)
        if result["success"]:
            return result["text"]
        return f"Error: {result.get('error', 'Unknown error')}"

    @function_tool
    async def extract_pdf_tables(file_path: str, page_number: int | None = None) -> str:
        """
        Extract structured table data from PDF files for analysis and processing.

        This tool identifies and extracts tables from PDF documents, converting them into
        structured data that can be easily analyzed or exported. You can extract tables
        from all pages or target a specific page.

        Parameters:
            file_path (str):
                The path to the PDF file containing tables.
                Example: "/path/to/financial_report.pdf"
            page_number (int | None):
                Optional. The specific page number to extract tables from (1-indexed).
                If None, extracts tables from all pages. Default: None.

        Returns:
            str:
                JSON string containing an array of tables. Each table includes:
                - page: The page number where the table was found
                - table_index: The index of the table on that page
                - data: 2D array representing the table structure (rows and columns)

        Usage:
            Use this tool when you need to extract tabular data from PDF documents like
            financial reports, data sheets, or research papers with tables. The extracted
            data can be further processed, analyzed, or converted to spreadsheets.
        """
        result = await _extract_pdf_tables(file_path, page_number)
        return json.dumps(result, indent=2)

    @function_tool
    async def merge_pdfs(
        ctx: RunContextWrapper[AgentExecutionContext],
        pdf_paths: list[str],
        output_path: str,
    ) -> str:
        """
        Combine multiple PDF files into a single unified PDF document.

        This tool takes multiple PDF files and merges them in the order provided,
        creating one consolidated PDF. Useful for combining reports, chapters, or
        related documents into a single file.

        **✨ AUTOMATIC FILE TRACKING:**
        Files you create are **automatically tracked as resources**!

        Parameters:
            pdf_paths (list[str]):
                List of paths to PDF files to merge, in the desired order.
                Example: ["/path/to/chapter1.pdf", "/path/to/chapter2.pdf"]
            output_path (str):
                The path where the merged PDF should be saved.
                Example: "/path/to/merged_document.pdf"

        Returns:
            str:
                Human-readable summary: "✅ Merged X PDFs into: filename.pdf"

        Usage:
            Call this tool when you need to combine multiple PDF documents into one file.
            Common use cases include: merging report sections, combining invoices,
            consolidating project documentation, or creating a single file from multiple
            scanned pages.
        """
        result = await _merge_pdfs(pdf_paths, output_path)

        # AUTO-REGISTER: Track created file as intermediary resource
        if result["success"]:
            from pathlib import Path

            try:
                if ctx.context:
                    from manager_agent_gym.schemas.domain.resource import Resource

                    filename = Path(result["output_path"]).name
                    file_size = Path(result["output_path"]).stat().st_size
                    ctx.context.register_created_resource(
                        Resource(
                            name=f"Generated: {filename}",
                            description="Auto-created by merge_pdfs tool",
                            file_path=result["output_path"],
                            mime_type="application/pdf",
                            size_bytes=file_size,
                            resource_role="intermediary",
                        )
                    )
            except Exception as e:
                from manager_agent_gym.core.common.logging import logger

                logger.warning(f"Failed to auto-register merged PDF resource: {e}")

            return f"✅ Merged {result['merged_count']} PDFs into: {Path(result['output_path']).name}\nLocation: {result['output_path']}"

        return f"❌ Error merging PDFs: {result.get('error', 'Unknown error')}"

    @function_tool
    async def create_simple_pdf(
        ctx: RunContextWrapper[AgentExecutionContext],
        text_content: str,
        output_path: str,
        title: str | None = None,
    ) -> str:
        """
        Generate a formatted PDF document from plain text content.

        This tool creates a professional-looking PDF from text content, with optional
        title formatting. The text is automatically formatted with proper line spacing,
        margins, and pagination. Paragraphs separated by double newlines are preserved.

        **✨ AUTOMATIC FILE TRACKING:**
        Files you create are **automatically tracked as resources**!

        Parameters:
            text_content (str):
                The text content to convert into a PDF. Use double newlines (\n\n) to
                separate paragraphs.
            output_path (str):
                The path where the PDF should be saved.
                Example: "/path/to/output.pdf"
            title (str | None):
                Optional. A title to display at the top of the document. Default: None.

        Returns:
            str:
                Human-readable summary: "✅ Created PDF: filename.pdf"

        Usage:
            Use this tool to create PDF documents from generated text, reports, summaries,
            or any text content that needs to be shared in PDF format. Perfect for creating
            formal documents, reports, or letters.
        """
        result = await _create_pdf(text_content, output_path, title)

        # AUTO-REGISTER: Track created file as intermediary resource
        if result["success"]:
            from pathlib import Path

            try:
                if ctx.context:
                    from manager_agent_gym.schemas.domain.resource import Resource

                    filename = Path(result["output_path"]).name
                    file_size = Path(result["output_path"]).stat().st_size
                    ctx.context.register_created_resource(
                        Resource(
                            name=f"Generated: {filename}",
                            description="Auto-created by create_simple_pdf tool",
                            file_path=result["output_path"],
                            mime_type="application/pdf",
                            size_bytes=file_size,
                            resource_role="intermediary",
                        )
                    )
            except Exception as e:
                from manager_agent_gym.core.common.logging import logger

                logger.warning(f"Failed to auto-register PDF resource: {e}")

            return f"✅ Created PDF: {Path(result['output_path']).name}\nLocation: {result['output_path']}"

        return f"❌ Error creating PDF: {result.get('error', 'Unknown error')}"

    @function_tool
    async def read_docx(file_path: str) -> str:
        """
        Extract text content from Microsoft Word (.docx) documents.

        This tool reads Word documents and extracts all text content including paragraphs
        and tables. Table data is formatted with pipe separators (|) for clarity. Perfect
        for analyzing contracts, reports, proposals, or any Word documents.

        Parameters:
            file_path (str):
                The path to the Word document (.docx file) to read.
                Example: "/path/to/document.docx" or "reports/project_proposal.docx"

        Returns:
            str:
                The full text content of the Word document with paragraphs separated by
                double newlines, or an error message if the file cannot be read.

        Usage:
            Call this tool when you need to read and analyze Word documents. Use it for
            extracting information, reviewing document content, summarizing reports, or
            searching for specific data within .docx files.
        """
        result = await _read_docx(file_path)
        if result["success"]:
            return result["text"]
        return f"Error: {result.get('error', 'Unknown error')}"

    @function_tool
    async def create_docx(
        ctx: RunContextWrapper[AgentExecutionContext],
        content: str,
        output_path: str,
        title: str | None = None,
        template_style: str = "normal",
    ) -> str:
        """
        Generate a professionally formatted Microsoft Word document from text content.

        This tool creates Word documents with customizable formatting styles. It supports
        Markdown-style headings (# for Heading 1, ## for Heading 2, etc.) and offers
        different template styles for various document types like formal letters, memos,
        or standard documents.

        **✨ AUTOMATIC FILE TRACKING:**
        Files you create are **automatically tracked as resources**! Just create the file
        and reference it in your notes - the system handles resource management.

        Parameters:
            content (str):
                The text content for the document. Use double newlines (\n\n) to separate
                paragraphs. Markdown-style headings (# Heading) are automatically converted
                to Word heading styles.
            output_path (str):
                The path where the .docx file should be saved.
                Example: "/path/to/output.docx"
            title (str | None):
                Optional. A centered title displayed at the top of the document. Default: None.
            template_style (str):
                The formatting template to use. Options: "normal" (default), "formal"
                (1.25" margins, 12pt font), or "memo" (smaller margins, 11pt font).

        Returns:
            str:
                Human-readable summary: "✅ Created Word document: filename.docx (X KB)"

        Usage:
            Use this tool to create professional Word documents for reports, letters,
            proposals, or documentation. The template styles help match different business
            communication needs.
        """
        result = await _create_docx(content, output_path, title, template_style)

        # AUTO-REGISTER: Track created file as intermediary resource
        if result["success"]:
            from pathlib import Path

            try:
                if ctx.context:
                    from manager_agent_gym.schemas.domain.resource import Resource

                    filename = Path(result["output_path"]).name
                    ctx.context.register_created_resource(
                        Resource(
                            name=f"Generated: {filename}",
                            description="Auto-created by create_docx tool",
                            file_path=result["output_path"],
                            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            size_bytes=result["file_size"],
                            resource_role="intermediary",
                        )
                    )
            except Exception as e:
                from manager_agent_gym.core.common.logging import logger

                logger.warning(f"Failed to auto-register DOCX resource: {e}")

            return f"✅ Created Word document: {Path(result['output_path']).name} ({result['file_size']} bytes)\nLocation: {result['output_path']}"

        return (
            f"❌ Error creating Word document: {result.get('error', 'Unknown error')}"
        )

    @function_tool
    async def add_table_to_docx(
        file_path: str, table_data: list[list[str]], header_row: bool = True
    ) -> str:
        """
        Insert a formatted table into an existing Word document.

        This tool appends a table to an existing Word document with automatic styling
        and formatting. The first row can be formatted as a header with bold text.
        Perfect for adding data tables, comparison charts, or structured information
        to reports.

        Parameters:
            file_path (str):
                The path to an existing Word document where the table will be added.
                Example: "/path/to/report.docx"
            table_data (list[list[str]]):
                A 2D array representing the table. Each inner list is a row, and elements
                are cell values. Example: [["Name", "Age"], ["Alice", "30"], ["Bob", "25"]]
            header_row (bool):
                Whether to format the first row as a header (bold, styled). Default: True.

        Returns:
            str:
                Success message confirming the table was added, or an error message if
                the operation fails.

        Usage:
            Call this tool when you need to add tabular data to an existing Word document.
            Common uses include adding data summaries, comparison tables, schedules, or
            any structured information to reports or documents.
        """
        result = await _add_table_to_docx(file_path, table_data, header_row)
        if result["success"]:
            return f"Successfully added table to {file_path}"
        return f"Error: {result.get('error', 'Unknown error')}"

    @function_tool
    async def extract_docx_structure(file_path: str) -> str:
        """
        Analyze and extract the structural outline of a Word document.

        This tool analyzes a Word document and returns its structural components without
        the full text. It identifies headings, paragraphs, and tables with their styles
        and positions. Useful for understanding document organization or creating a
        document outline.

        Parameters:
            file_path (str):
                The path to the Word document to analyze.
                Example: "/path/to/document.docx"

        Returns:
            str:
                JSON string containing the document structure. Each element includes:
                - type: "paragraph" or "table"
                - style: The Word style name (e.g., "Heading 1", "Normal")
                - text: Preview of paragraph text (first 100 characters) or table dimensions
                - rows/columns: For tables, the number of rows and columns

        Usage:
            Use this tool when you need to understand a document's organization without
            reading the full content. Perfect for quickly assessing document structure,
            finding specific sections, or understanding how a document is organized before
            extracting specific parts.
        """
        result = await _extract_docx_structure(file_path)
        return json.dumps(result, indent=2)

    @function_tool
    async def convert_docx_to_pdf(
        ctx: RunContextWrapper[AgentExecutionContext], docx_path: str, output_path: str
    ) -> str:
        """
        Convert a Microsoft Word document to PDF format while preserving formatting.

        This tool converts .docx or .doc files to PDF format using LibreOffice, preserving all
        formatting, styles, images, and layout. The resulting PDF maintains the
        professional appearance of the original Word document.

        **✨ AUTOMATIC FILE TRACKING:**
        Files you create are **automatically tracked as resources**!

        IMPORTANT: This tool ONLY accepts Word documents (.docx, .doc files).
        If you have Markdown (.md) or other formats, convert them to DOCX first using
        'convert_markdown_to_docx' before calling this tool.

        Note: Requires LibreOffice to be installed on the system.

        Parameters:
            docx_path (str):
                The path to the Word document to convert. Must be a .docx or .doc file.
                Example: "/path/to/document.docx"
            output_path (str):
                The path where the PDF should be saved.
                Example: "/path/to/document.pdf"

        Returns:
            str:
                Human-readable summary: "✅ Converted to PDF: filename.pdf"

        Usage:
            Call this tool when you need to convert Word documents to PDF for sharing,
            archiving, or distribution. PDFs ensure formatting consistency across
            different systems and prevent editing.

            For Markdown files: First use 'convert_markdown_to_docx', then use this tool.
        """
        result = await _convert_docx_to_pdf(docx_path, output_path)

        # AUTO-REGISTER: Track created file as intermediary resource
        if result["success"]:
            from pathlib import Path

            try:
                if ctx.context:
                    from manager_agent_gym.schemas.domain.resource import Resource

                    filename = Path(result["output_path"]).name
                    file_size = Path(result["output_path"]).stat().st_size
                    ctx.context.register_created_resource(
                        Resource(
                            name=f"Generated: {filename}",
                            description="Auto-created by convert_docx_to_pdf tool",
                            file_path=result["output_path"],
                            mime_type="application/pdf",
                            size_bytes=file_size,
                            resource_role="intermediary",
                        )
                    )
            except Exception as e:
                from manager_agent_gym.core.common.logging import logger

                logger.warning(f"Failed to auto-register converted PDF resource: {e}")

            return f"✅ Converted to PDF: {Path(result['output_path']).name}\nLocation: {result['output_path']}"

        return f"❌ Error converting to PDF: {result.get('error', 'Unknown error')}"

    @function_tool
    async def convert_markdown_to_docx(
        ctx: RunContextWrapper[AgentExecutionContext],
        markdown_content: str,
        output_path: str,
    ) -> str:
        """
        Convert Markdown-formatted text into a professionally styled Word document.

        This tool parses Markdown syntax and converts it to proper Word document formatting.
        Headings, lists, code blocks, and paragraphs are all converted to appropriate
        Word styles. Perfect for turning documentation, notes, or Markdown files into
        editable Word documents.

        **✨ AUTOMATIC FILE TRACKING:**
        Files you create are **automatically tracked as resources**!

        Parameters:
            markdown_content (str):
                The Markdown text to convert. Supports:
                - Headings: # H1, ## H2, ### H3
                - Lists: - bullet or * bullet, and numbered lists (1. item)
                - Code blocks: ```code```
                - Paragraphs: separated by blank lines
            output_path (str):
                The path where the .docx file should be saved.
                Example: "/path/to/output.docx"

        Returns:
            str:
                Human-readable summary: "✅ Converted Markdown to DOCX: filename.docx"

        Usage:
            Use this tool to convert Markdown content (from files, generated text, or
            documentation) into Word documents. Ideal for creating editable versions of
            Markdown documentation or converting notes into formal documents.
        """
        result = await _convert_markdown_to_docx(markdown_content, output_path)

        # AUTO-REGISTER: Track created file as intermediary resource
        if result["success"]:
            from pathlib import Path

            try:
                if ctx.context:
                    from manager_agent_gym.schemas.domain.resource import Resource

                    filename = Path(result["output_path"]).name
                    file_size = Path(result["output_path"]).stat().st_size
                    ctx.context.register_created_resource(
                        Resource(
                            name=f"Generated: {filename}",
                            description="Auto-created by convert_markdown_to_docx tool",
                            file_path=result["output_path"],
                            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            size_bytes=file_size,
                            resource_role="intermediary",
                        )
                    )
            except Exception as e:
                from manager_agent_gym.core.common.logging import logger

                logger.warning(f"Failed to auto-register converted DOCX resource: {e}")

            return f"✅ Converted Markdown to DOCX: {Path(result['output_path']).name}\nLocation: {result['output_path']}"

        return f"❌ Error converting Markdown to DOCX: {result.get('error', 'Unknown error')}"

    @function_tool
    async def convert_text_to_docx(
        ctx: RunContextWrapper[AgentExecutionContext],
        text_content: str,
        output_path: str,
        font_size: int = 11,
    ) -> str:
        """
        Convert plain text into a Word document with customizable font size.

        This tool takes plain text and creates a simple Word document with consistent
        formatting. Paragraphs separated by double newlines are preserved. Useful for
        converting plain text files or generated content into editable Word documents.

        **✨ AUTOMATIC FILE TRACKING:**
        Files you create are **automatically tracked as resources**!

        Parameters:
            text_content (str):
                The plain text content to convert. Use double newlines (\n\n) to
                separate paragraphs.
            output_path (str):
                The path where the .docx file should be saved.
                Example: "/path/to/output.docx"
            font_size (int):
                The font size in points for the document text. Default: 11.

        Returns:
            str:
                Human-readable summary: "✅ Converted text to DOCX: filename.docx"

        Usage:
            Use this tool to convert plain text content into editable Word documents.
            Perfect for converting text files, logs, or generated text into a format
            that can be easily formatted and shared.
        """
        result = await _convert_text_to_docx(text_content, output_path, font_size)

        # AUTO-REGISTER: Track created file as intermediary resource
        if result["success"]:
            from pathlib import Path

            try:
                if ctx.context:
                    from manager_agent_gym.schemas.domain.resource import Resource

                    filename = Path(result["output_path"]).name
                    file_size = Path(result["output_path"]).stat().st_size
                    ctx.context.register_created_resource(
                        Resource(
                            name=f"Generated: {filename}",
                            description="Auto-created by convert_text_to_docx tool",
                            file_path=result["output_path"],
                            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            size_bytes=file_size,
                            resource_role="intermediary",
                        )
                    )
            except Exception as e:
                from manager_agent_gym.core.common.logging import logger

                logger.warning(f"Failed to auto-register converted DOCX resource: {e}")

            return f"✅ Converted text to DOCX: {Path(result['output_path']).name}\nLocation: {result['output_path']}"

        return (
            f"❌ Error converting text to DOCX: {result.get('error', 'Unknown error')}"
        )

    @function_tool
    async def save_text_as_markdown(
        ctx: RunContextWrapper[AgentExecutionContext], content: str, file_name: str
    ) -> str:
        """
        Save text content as a Markdown (.md) file with automatic formatting detection.

        This tool saves text content as a Markdown file, automatically appending the .md
        extension if needed. Perfect for saving documentation, notes, or formatted text.

        **✨ AUTOMATIC FILE TRACKING:**
        Files you create are **automatically tracked as resources**! Just create the file
        and reference it in your notes - the system handles resource management.

        Parameters:
            content (str):
                The text content to save. Can include Markdown formatting like headings (#),
                lists, code blocks, etc.
            file_name (str):
                The name for the file. The .md extension is added automatically if not present.
                Example: "notes" or "documentation.md"

        Returns:
            str:
                Human-readable summary: "✅ Saved markdown: filename.md (X words)"

        Usage:
            Use this tool to save generated or processed text as Markdown files. Ideal for
            creating documentation, saving notes, or storing formatted text content that
            might be used later or shared with others.
        """
        result = await _save_markdown(content, file_name)

        # AUTO-REGISTER: Track created file as intermediary resource
        if result["success"]:
            try:
                if ctx.context:
                    from manager_agent_gym.schemas.domain.resource import Resource

                    ctx.context.register_created_resource(
                        Resource(
                            name=f"Generated: {result['file_name']}",
                            description="Auto-created by save_text_as_markdown tool",
                            file_path=result["file_path"],
                            mime_type=result["mime_type"],
                            size_bytes=result["size_bytes"],
                            resource_role="intermediary",
                        )
                    )
            except Exception as e:
                from manager_agent_gym.core.common.logging import logger

                logger.warning(f"Failed to auto-register markdown resource: {e}")

            word_count = result["content_summary"]["word_count"]
            return (
                f"✅ Saved markdown: {result['file_name']} "
                f"({word_count} words, {result['size_bytes']} bytes)\n"
                f"Location: {result['file_path']}"
            )

        return f"❌ Error saving markdown: {result.get('error', 'Unknown error')}"

    return [
        read_pdf,
        extract_pdf_tables,
        merge_pdfs,
        create_simple_pdf,
        read_docx,
        create_docx,
        add_table_to_docx,
        extract_docx_structure,
        convert_docx_to_pdf,
        convert_markdown_to_docx,
        convert_text_to_docx,
        save_text_as_markdown,
    ]
